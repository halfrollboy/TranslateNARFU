from docx import Document
import csv 

class Source:
    def __init__(self, file_name):
        self.file_name = file_name
        self.text = ""
        self.lines = []

    def parse(self):
        raise NotImplementedError
    
    def save_path(self, mb_text):
        raise NotImplementedError


class TxtSource(Source):
    def parse(self):
        try:
            with open(self.file_name) as buf:
                self.text = buf.read()
            self.lines = self.text.split('\n')
        except Exception as ex: # TODO: too common exception
            return False
        return True

    def save_path(self, mb_text):
        txtfile = open(self.file_name[0]+self.file_name[1], 'w')
        txtfile.write(mb_text)
        txtfile.close()


class DocxSource(Source):
    def parse(self):
        # TODO: stub
        try:
            doc=Document(self.file_name)
            for para in doc.paragraphs:
                self.lines.append(para.text)
            self.text ='\n'.join(self.lines)
        except Exception as ex: # TODO: too common exception
            return False
        return True

    def save_path(self, mb_text): #сохраняет но только в дерикторию с проектом
        document = Document()            
        document.add_heading('Приветственный title', 0)

        p = document.add_paragraph(self.text)
        document.save(self.file_name[0]+self.file_name[1])


class CSVSource(Source):
    def __init__(self, file_name):
        super().__init__(file_name)
        self.left_word=''
        self.possible_phrases=[]
        self.any_word={}
        self._append=self.lines.append

    def find_word_in_csv_dict(self,word):
        #l=','
        p=';' #костыли - я не знаю почему они в файле появляются 
        with open(self.file_name,'r',newline="") as csv_f:
            csv_data=csv.reader(csv_f)
            for row in csv_data:
                t=','.join(row).replace(p,'')
                if t.find( word )>=0:
                    self.possible_phrases.append(t) #TODO переделать поиск 

    #Функция для сзодания списка терминов split метод
    def termin_list(self):
        self.lines.clear()
        with open(self.file_name,'r') as f:
            csv_data=csv.reader(f)
            for row in csv_data:
                line = row.split()
                for i in line:
                    if i.isupper():
                        perem+=i+" "
                        self._append(perem)

    #Получение длинны словаря 
    def get_len_dict(self):# TODO добавить корректность
        reader = list(csv.DictReader(open(self.file_name), delimiter='|'))
        return(len(reader))

    #Добавление нового слова
    def set_new_word(self, new_word, description):
        with open(self.file_name,'a') as f:
            writer=csv.writer(f)
            writer.writerow((new_word,description))

    #Перезапись словаря в класс CSV
    def get_dict_in_this_class(self):
        self.lines.clear()
        p=';'
        with open(self.file_name,'r') as f:
            csv_data=csv.reader(f)
            for row in csv_data:
                t=','.join(row).replace(p,'')
                t.find(',')
                self._append(t)

    #Запись словаря в основной класс программы 
    def get_dict(self):
        self.lines.clear()
        p=';'
        with open(self.file_name,'r') as f:
            csv_data=csv.reader(f)
            for row in csv_data:
                t=','.join(row).replace(p,'')
                t.find(',')
                self._append(t)
        return self.lines

    #Перезапись словаря терминов
    def get_dict_terms(self):
        self.lines.clear()
        p=';'
        with open(self.file_name,'r') as f:
            csv_data=csv.reader(f)
            for row in csv_data:
                t=','.join(row).replace(p,'')
                termin=''
                for j in t:
                    if j!=',':
                        termin+=j
                    else:
                        self._append(termin)
                        break 
                            #ANTHEMORRHAGIC

    #--Проверка слова через словарь класса 
    #0)Составояем словарь для слова
    #1)Ищем старое слово в нововом списке 
    #2)Выписываем в possible_phrases=[] найденные совпадения 
    #3)Записываем слово в left_word
    #4)Приходит новое слово
    #5)Проверяем на наличие сходства в буфере, выводим старое слово и список совпадений из буфера
    def check_dict_for_word(self,word):
        word=word.upper()
        buff_list=[]
        buff_for_findings=[]

    #Ищем слово в словаре
        for i in self.lines:
            try:
                if word in i:
                    buff_list.append(i)
            except Exception:
                pass
        print('_________________',buff_list)
        # if len(buff_list)==0:
        #     pass

    #Ищем слово в облаке
        for j in self.possible_phrases:
            #print('j в поиске слова в облаке',j)
            try:
                if word in j:
                    buff_for_findings.append(j)
            except Exception:
                pass
        print('длинна списка вывода',len(buff_for_findings))
        if len(buff_for_findings)==0:
            buff_for_findings.extend(self.possible_phrases)

    #Проерка на вхождение слова из left_word
        print('buff_list ',buff_list)
        print('buff_finder ',buff_for_findings)
        print('possible_phrases ',self.possible_phrases)
        self.possible_phrases.clear()    
        for k in buff_list:
            try:
                if self.left_word in k:
                    self.possible_phrases.append(k)
            except Exception:
                pass
        if len(self.possible_phrases)==0:
            self.possible_phrases.extend(buff_list)

        print('possible_phrases|| 2',self.possible_phrases)

        self.left_word=word
        return buff_for_findings
#сделать чтобы выбирало только либо одно слово либо со встречающимися допами 

    def find_in_dict(self, word):
        self.possible_phrases.clear()
        for i in self.lines:
            try:
                if word in i:
                    self.possible_phrases.append(i)
            except Exception:
                pass
        if len(self.possible_phrases)==0 or len(self.possible_phrases)==None:
            return 
        if len(self.any_word)>4:
            self.any_word.pop(0)


SOURCES = {
    '.txt': TxtSource,
    '.docx': DocxSource,
}

